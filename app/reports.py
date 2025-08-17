from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import List, Dict, Any
from datetime import datetime, date
import os
from app.database import get_db
from app.models import SARCase, CaseUpdate, CaseFile, ICOEscalation
from app.crud import get_sar_case_db, get_case_updates_db, get_case_files_db
from io import BytesIO

def generate_pdf_report(sar_id: int, user_id: int) -> str:
    """Generate a comprehensive PDF report for a SAR case."""
    # Get case data
    sar_case = get_sar_case_db(sar_id, user_id)
    if not sar_case:
        raise ValueError("SAR case not found")
    
    updates = get_case_updates_db(sar_id, user_id)
    files = get_case_files_db(sar_id, user_id)
    
    # Create PDF document
    filename = f"SAR_Report_{sar_case.case_reference}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = os.path.join("uploads", "reports", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors.darkblue
    )
    
    # Build story
    story = []
    
    # Title page
    story.append(Paragraph("Subject Access Request Report", title_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Case Reference: {sar_case.case_reference}", styles['Heading3']))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
    story.append(PageBreak())
    
    # Case Overview
    story.append(Paragraph("Case Overview", heading_style))
    
    overview_data = [
        ["Case Reference", sar_case.case_reference],
        ["Organization", sar_case.organization_name],
        ["Request Type", sar_case.request_type],
        ["Submission Date", sar_case.submission_date.strftime("%d/%m/%Y")],
        ["Submission Method", sar_case.submission_method],
        ["Status", sar_case.status],
        ["Statutory Deadline", sar_case.statutory_deadline.strftime("%d/%m/%Y")],
    ]
    
    if sar_case.extended_deadline:
        overview_data.append(["Extended Deadline", sar_case.extended_deadline.strftime("%d/%m/%Y")])
    if sar_case.custom_deadline:
        overview_data.append(["Custom Deadline", sar_case.custom_deadline.strftime("%d/%m/%Y")])
    
    overview_table = Table(overview_data, colWidths=[2*inch, 4*inch])
    overview_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(overview_table)
    story.append(Spacer(1, 20))
    
    # Request Description
    story.append(Paragraph("Request Description", heading_style))
    story.append(Paragraph(sar_case.request_description, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Organization Details
    story.append(Paragraph("Organization Details", heading_style))
    
    org_data = [
        ["Name", sar_case.organization_name],
        ["Address", sar_case.organization_address or "Not provided"],
        ["Email", sar_case.organization_email or "Not provided"],
        ["Phone", sar_case.organization_phone or "Not provided"]
    ]
    
    org_table = Table(org_data, colWidths=[2*inch, 4*inch])
    org_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(org_table)
    story.append(Spacer(1, 20))
    
    # Timeline of Events
    story.append(Paragraph("Timeline of Events", heading_style))
    
    timeline_data = [["Date", "Event", "Details"]]
    
    # Add submission
    timeline_data.append([
        sar_case.submission_date.strftime("%d/%m/%Y"),
        "SAR Submitted",
        f"Request submitted via {sar_case.submission_method}"
    ])
    
    # Add updates
    for update in updates:
        timeline_data.append([
            update.created_at.strftime("%d/%m/%Y"),
            update.update_type,
            update.title
        ])
    
    # Add response if received
    if sar_case.response_received and sar_case.response_date:
        timeline_data.append([
            sar_case.response_date.strftime("%d/%m/%Y"),
            "Response Received",
            f"Response received from {sar_case.organization_name}"
        ])
    
    timeline_table = Table(timeline_data, colWidths=[1.5*inch, 2*inch, 2.5*inch])
    timeline_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(timeline_table)
    story.append(Spacer(1, 20))
    
    # Case Updates
    if updates:
        story.append(Paragraph("Case Updates", heading_style))
        
        for update in updates:
            story.append(Paragraph(f"<b>{update.created_at.strftime('%d/%m/%Y')} - {update.title}</b>", styles['Heading4']))
            story.append(Paragraph(f"<b>Type:</b> {update.update_type}", styles['Normal']))
            story.append(Paragraph(update.content, styles['Normal']))
            
            if update.correspondence_date:
                story.append(Paragraph(f"<b>Correspondence Date:</b> {update.correspondence_date.strftime('%d/%m/%Y')}", styles['Normal']))
            if update.correspondence_method:
                story.append(Paragraph(f"<b>Method:</b> {update.correspondence_method}", styles['Normal']))
            if update.call_duration:
                story.append(Paragraph(f"<b>Call Duration:</b> {update.call_duration} minutes", styles['Normal']))
            
            story.append(Spacer(1, 12))
    
    # Attached Files
    if files:
        story.append(Paragraph("Attached Files", heading_style))
        
        files_data = [["Filename", "Category", "Size", "Uploaded"]]
        
        for file in files:
            files_data.append([
                file.original_filename,
                file.file_category,
                f"{file.file_size / 1024:.1f} KB",
                file.uploaded_at.strftime("%d/%m/%Y")
            ])
        
        files_table = Table(files_data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 1*inch])
        files_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(files_table)
        story.append(Spacer(1, 20))
    
    # ICO Escalations
    ico_escalations = [e for e in sar_case.ico_escalations]
    if ico_escalations:
        story.append(Paragraph("ICO Escalations", heading_style))
        
        for escalation in ico_escalations:
            story.append(Paragraph(f"<b>ICO Reference: {escalation.ico_reference}</b>", styles['Heading4']))
            story.append(Paragraph(f"<b>Escalation Date:</b> {escalation.escalation_date.strftime('%d/%m/%Y')}", styles['Normal']))
            story.append(Paragraph(f"<b>Reason:</b> {escalation.escalation_reason}", styles['Normal']))
            story.append(Paragraph(f"<b>Status:</b> {escalation.status}", styles['Normal']))
            
            if escalation.ico_decision:
                story.append(Paragraph(f"<b>ICO Decision:</b> {escalation.ico_decision}", styles['Normal']))
                story.append(Paragraph(f"<b>Decision Date:</b> {escalation.ico_decision_date.strftime('%d/%m/%Y')}", styles['Normal']))
                story.append(Paragraph(f"<b>Summary:</b> {escalation.ico_decision_summary}", styles['Normal']))
            
            story.append(Spacer(1, 12))
    
    # Compliance Assessment
    story.append(Paragraph("Compliance Assessment", heading_style))
    
    current_date = date.today()
    deadline = sar_case.extended_deadline or sar_case.custom_deadline or sar_case.statutory_deadline
    days_overdue = (current_date - deadline).days if current_date > deadline else 0
    
    if sar_case.response_received:
        response_time = (sar_case.response_date - sar_case.submission_date).days
        compliance_status = "Compliant" if response_time <= 28 else "Non-compliant"
        story.append(Paragraph(f"<b>Response Status:</b> {compliance_status}", styles['Normal']))
        story.append(Paragraph(f"<b>Response Time:</b> {response_time} days", styles['Normal']))
        story.append(Paragraph(f"<b>Response Date:</b> {sar_case.response_date.strftime('%d/%m/%Y')}", styles['Normal']))
    else:
        if days_overdue > 0:
            story.append(Paragraph(f"<b>Status:</b> Non-compliant - {days_overdue} days overdue", styles['Normal']))
        else:
            days_remaining = (deadline - current_date).days
            story.append(Paragraph(f"<b>Status:</b> Pending - {days_remaining} days remaining", styles['Normal']))
    
    story.append(Spacer(1, 20))
    
    # Recommendations
    story.append(Paragraph("Recommendations", heading_style))
    
    if sar_case.status == "Overdue":
        story.append(Paragraph("• Send immediate follow-up communication", styles['Normal']))
        story.append(Paragraph("• Consider ICO escalation if no response within 7 days", styles['Normal']))
        story.append(Paragraph("• Document all communication attempts", styles['Normal']))
    elif sar_case.status == "Pending":
        story.append(Paragraph("• Monitor deadline approaching", styles['Normal']))
        story.append(Paragraph("• Prepare follow-up communication", styles['Normal']))
    elif sar_case.status == "Responded":
        story.append(Paragraph("• Review response for completeness", styles['Normal']))
        story.append(Paragraph("• Update case status", styles['Normal']))
        if not sar_case.data_complete:
            story.append(Paragraph("• Consider follow-up request for missing data", styles['Normal']))
    
    # Build PDF
    doc.build(story)
    
    return filepath

def generate_word_report(sar_id: int, user_id: int) -> str:
    """Generate a comprehensive Word document report for a SAR case."""
    # Get case data
    sar_case = get_sar_case_db(sar_id, user_id)
    if not sar_case:
        raise ValueError("SAR case not found")
    
    updates = get_case_updates_db(sar_id, user_id)
    files = get_case_files_db(sar_id, user_id)
    
    # Create Word document
    filename = f"SAR_Report_{sar_case.case_reference}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join("uploads", "reports", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    doc = Document()
    
    # Title
    title = doc.add_heading("Subject Access Request Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Case Reference
    doc.add_paragraph(f"Case Reference: {sar_case.case_reference}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_page_break()
    
    # Case Overview
    doc.add_heading("Case Overview", level=1)
    
    overview_table = doc.add_table(rows=1, cols=2)
    overview_table.style = 'Table Grid'
    hdr_cells = overview_table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    overview_data = [
        ["Case Reference", sar_case.case_reference],
        ["Organization", sar_case.organization_name],
        ["Request Type", sar_case.request_type],
        ["Submission Date", sar_case.submission_date.strftime("%d/%m/%Y")],
        ["Submission Method", sar_case.submission_method],
        ["Status", sar_case.status],
        ["Statutory Deadline", sar_case.statutory_deadline.strftime("%d/%m/%Y")],
    ]
    
    if sar_case.extended_deadline:
        overview_data.append(["Extended Deadline", sar_case.extended_deadline.strftime("%d/%m/%Y")])
    if sar_case.custom_deadline:
        overview_data.append(["Custom Deadline", sar_case.custom_deadline.strftime("%d/%m/%Y")])
    
    for field, value in overview_data:
        row_cells = overview_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = str(value)
    
    # Request Description
    doc.add_heading("Request Description", level=1)
    doc.add_paragraph(sar_case.request_description)
    
    # Organization Details
    doc.add_heading("Organization Details", level=1)
    
    org_table = doc.add_table(rows=1, cols=2)
    org_table.style = 'Table Grid'
    hdr_cells = org_table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    org_data = [
        ["Name", sar_case.organization_name],
        ["Address", sar_case.organization_address or "Not provided"],
        ["Email", sar_case.organization_email or "Not provided"],
        ["Phone", sar_case.organization_phone or "Not provided"]
    ]
    
    for field, value in org_data:
        row_cells = org_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = str(value)
    
    # Timeline of Events
    doc.add_heading("Timeline of Events", level=1)
    
    timeline_table = doc.add_table(rows=1, cols=3)
    timeline_table.style = 'Table Grid'
    hdr_cells = timeline_table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Event'
    hdr_cells[2].text = 'Details'
    
    # Add submission
    row_cells = timeline_table.add_row().cells
    row_cells[0].text = sar_case.submission_date.strftime("%d/%m/%Y")
    row_cells[1].text = "SAR Submitted"
    row_cells[2].text = f"Request submitted via {sar_case.submission_method}"
    
    # Add updates
    for update in updates:
        row_cells = timeline_table.add_row().cells
        row_cells[0].text = update.created_at.strftime("%d/%m/%Y")
        row_cells[1].text = update.update_type
        row_cells[2].text = update.title
    
    # Add response if received
    if sar_case.response_received and sar_case.response_date:
        row_cells = timeline_table.add_row().cells
        row_cells[0].text = sar_case.response_date.strftime("%d/%m/%Y")
        row_cells[1].text = "Response Received"
        row_cells[2].text = f"Response received from {sar_case.organization_name}"
    
    # Case Updates
    if updates:
        doc.add_heading("Case Updates", level=1)
        
        for update in updates:
            doc.add_heading(f"{update.created_at.strftime('%d/%m/%Y')} - {update.title}", level=2)
            doc.add_paragraph(f"Type: {update.update_type}")
            doc.add_paragraph(update.content)
            
            if update.correspondence_date:
                doc.add_paragraph(f"Correspondence Date: {update.correspondence_date.strftime('%d/%m/%Y')}")
            if update.correspondence_method:
                doc.add_paragraph(f"Method: {update.correspondence_method}")
            if update.call_duration:
                doc.add_paragraph(f"Call Duration: {update.call_duration} minutes")
    
    # Attached Files
    if files:
        doc.add_heading("Attached Files", level=1)
        
        files_table = doc.add_table(rows=1, cols=4)
        files_table.style = 'Table Grid'
        hdr_cells = files_table.rows[0].cells
        hdr_cells[0].text = 'Filename'
        hdr_cells[1].text = 'Category'
        hdr_cells[2].text = 'Size'
        hdr_cells[3].text = 'Uploaded'
        
        for file in files:
            row_cells = files_table.add_row().cells
            row_cells[0].text = file.original_filename
            row_cells[1].text = file.file_category
            row_cells[2].text = f"{file.file_size / 1024:.1f} KB"
            row_cells[3].text = file.uploaded_at.strftime("%d/%m/%Y")
    
    # ICO Escalations
    ico_escalations = [e for e in sar_case.ico_escalations]
    if ico_escalations:
        doc.add_heading("ICO Escalations", level=1)
        
        for escalation in ico_escalations:
            doc.add_heading(f"ICO Reference: {escalation.ico_reference}", level=2)
            doc.add_paragraph(f"Escalation Date: {escalation.escalation_date.strftime('%d/%m/%Y')}")
            doc.add_paragraph(f"Reason: {escalation.escalation_reason}")
            doc.add_paragraph(f"Status: {escalation.status}")
            
            if escalation.ico_decision:
                doc.add_paragraph(f"ICO Decision: {escalation.ico_decision}")
                doc.add_paragraph(f"Decision Date: {escalation.ico_decision_date.strftime('%d/%m/%Y')}")
                doc.add_paragraph(f"Summary: {escalation.ico_decision_summary}")
    
    # Compliance Assessment
    doc.add_heading("Compliance Assessment", level=1)
    
    current_date = date.today()
    deadline = sar_case.extended_deadline or sar_case.custom_deadline or sar_case.statutory_deadline
    days_overdue = (current_date - deadline).days if current_date > deadline else 0
    
    if sar_case.response_received:
        response_time = (sar_case.response_date - sar_case.submission_date).days
        compliance_status = "Compliant" if response_time <= 28 else "Non-compliant"
        doc.add_paragraph(f"Response Status: {compliance_status}")
        doc.add_paragraph(f"Response Time: {response_time} days")
        doc.add_paragraph(f"Response Date: {sar_case.response_date.strftime('%d/%m/%Y')}")
    else:
        if days_overdue > 0:
            doc.add_paragraph(f"Status: Non-compliant - {days_overdue} days overdue")
        else:
            days_remaining = (deadline - current_date).days
            doc.add_paragraph(f"Status: Pending - {days_remaining} days remaining")
    
    # Recommendations
    doc.add_heading("Recommendations", level=1)
    
    if sar_case.status == "Overdue":
        doc.add_paragraph("• Send immediate follow-up communication")
        doc.add_paragraph("• Consider ICO escalation if no response within 7 days")
        doc.add_paragraph("• Document all communication attempts")
    elif sar_case.status == "Pending":
        doc.add_paragraph("• Monitor deadline approaching")
        doc.add_paragraph("• Prepare follow-up communication")
    elif sar_case.status == "Responded":
        doc.add_paragraph("• Review response for completeness")
        doc.add_paragraph("• Update case status")
        if not sar_case.data_complete:
            doc.add_paragraph("• Consider follow-up request for missing data")
    
    # Save document
    doc.save(filepath)
    
    return filepath

def generate_initial_sar_letter(sar_case) -> bytes:
    """Generate initial SAR request letter as PDF."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from io import BytesIO
    from datetime import datetime
    
    # Create PDF buffer
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    
    # Get user information
    user = None # No user_id passed, so user info is not available here
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1,  # Center
        textColor=colors.darkblue
    )
    
    normal_style = styles['Normal']
    normal_style.fontSize = 11
    normal_style.spaceAfter = 12
    
    # Header
    story.append(Paragraph("SUBJECT ACCESS REQUEST", title_style))
    story.append(Spacer(1, 20))
    
    # Date
    current_date = datetime.now().strftime("%B %d, %Y")
    story.append(Paragraph(f"Date: {current_date}", normal_style))
    story.append(Spacer(1, 20))
    
    # Organization Information
    story.append(Paragraph("To:", normal_style))
    story.append(Paragraph(f"{sar_case.organization_name}", normal_style))
    
    if sar_case.organization_address:
        story.append(Paragraph(f"{sar_case.organization_address}", normal_style))
    
    if sar_case.organization_email:
        story.append(Paragraph(f"Email: {sar_case.organization_email}", normal_style))
    
    if sar_case.organization_phone:
        story.append(Paragraph(f"Phone: {sar_case.organization_phone}", normal_style))
    
    story.append(Spacer(1, 20))
    
    # Data Protection Officer/Administrator
    if sar_case.data_administrator_name:
        story.append(Paragraph(f"Data Protection Officer: {sar_case.data_administrator_name}", normal_style))
    
    if sar_case.data_controller_name:
        story.append(Paragraph(f"Data Controller: {sar_case.data_controller_name}", normal_style))
    
    story.append(Spacer(1, 20))
    
    # Subject
    story.append(Paragraph("Subject: Subject Access Request", normal_style))
    story.append(Spacer(1, 20))
    
    # Main Content
    story.append(Paragraph("Dear Sir/Madam,", normal_style))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph(
        "I am writing to make a formal Subject Access Request (SAR) under the Data Protection Act 2018 and UK GDPR Article 15.",
        normal_style
    ))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph(
        "I hereby request that you provide me with the following information:",
        normal_style
    ))
    story.append(Spacer(1, 12))
    
    # Request Details
    story.append(Paragraph(f"<b>Request Type:</b> {sar_case.request_type}", normal_style))
    story.append(Paragraph(f"<b>Case Reference:</b> {sar_case.case_reference}", normal_style))
    story.append(Paragraph(f"<b>Submission Date:</b> {sar_case.submission_date.strftime('%B %d, %Y')}", normal_style))
    story.append(Spacer(1, 12))
    
    # Request Description
    if sar_case.request_description:
        story.append(Paragraph("<b>Specific Request Details:</b>", normal_style))
        story.append(Paragraph(sar_case.request_description, normal_style))
        story.append(Spacer(1, 12))
    
    # Information Requested
    story.append(Paragraph("<b>I specifically request the following information:</b>", normal_style))
    story.append(Spacer(1, 12))
    
    info_items = [
        "1. Confirmation that you are processing my personal data",
        "2. The legal basis for processing my personal data",
        "3. The categories of personal data you hold about me",
        "4. The source of the personal data",
        "5. The purposes for which the data is being processed",
        "6. Any third parties with whom you share my personal data",
        "7. The retention period for my personal data",
        "8. My rights under data protection law",
        "9. Any automated decision-making processes",
        "10. A copy of my personal data in a structured, commonly used format"
    ]
    
    for item in info_items:
        story.append(Paragraph(item, normal_style))
    
    story.append(Spacer(1, 20))
    
    # Legal Basis
    story.append(Paragraph(
        "This request is made under Article 15 of the UK GDPR and Section 45 of the Data Protection Act 2018. "
        "You are required to respond to this request within one calendar month of receipt.",
        normal_style
    ))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph(
        "If you require additional information to identify me or locate the relevant data, please contact me immediately. "
        "If you need to extend the response period, you must inform me within one month of receiving this request.",
        normal_style
    ))
    story.append(Spacer(1, 20))
    
    # Contact Information
    story.append(Paragraph("<b>My Contact Information:</b>", normal_style))
    user_name = "Data Subject" # No user_id, so no user info
    user_email = "To be provided" # No user_id, so no user info
    story.append(Paragraph(f"Name: {user_name}", normal_style))
    story.append(Paragraph(f"Email: {user_email}", normal_style))
    story.append(Spacer(1, 20))
    
    # Closing
    story.append(Paragraph(
        "I look forward to receiving your response within the statutory timeframe. "
        "If you have any questions about this request, please do not hesitate to contact me.",
        normal_style
    ))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Yours faithfully,", normal_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(user_name, normal_style))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"<i>This SAR was submitted on {sar_case.submission_date.strftime('%B %d, %Y')} "
        f"and is being tracked under reference {sar_case.case_reference}.</i>",
        ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9, textColor=colors.grey)
    ))
    
    # Build PDF
    doc.build(story)
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content
